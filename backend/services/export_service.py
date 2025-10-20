"""
Document Export Service

Handles exporting documentation in multiple formats:
- Markdown (.md)
- DOCX (.docx)
- PDF (.pdf)
"""

import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2

# WeasyPrint is optional (requires GTK libraries on Windows)
# If not available, PDF export will show an error message
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    print(f"⚠️ WeasyPrint not available: {e}")
    print("📝 PDF export will not work. Markdown and DOCX exports are still available.")
    WEASYPRINT_AVAILABLE = False
    HTML = None
    CSS = None


class ExportService:
    """Service for exporting documentation in various formats."""
    
    def __init__(self):
        """Initialize export service."""
        pass
    
    def export_markdown(self, project_name, documentation_content):
        """
        Export documentation as Markdown file.
        
        Args:
            project_name: Name of the project
            documentation_content: Markdown content string
            
        Returns:
            tuple: (file_content, filename)
        """
        # Add project header to markdown
        header = f"# {project_name}\n\n"
        header += f"*Generated on {datetime.now().strftime('%B %d, %Y')}*\n\n"
        header += "---\n\n"
        
        # Combine header with content
        full_content = header + documentation_content
        
        # Create filename
        filename = f"{self._sanitize_filename(project_name)}_documentation.md"
        
        return full_content.encode('utf-8'), filename
    
    def export_docx(self, project_name, documentation_content):
        """
        Export documentation as DOCX file.
        
        Args:
            project_name: Name of the project
            documentation_content: Markdown content string
            
        Returns:
            tuple: (file_bytes, filename)
        """
        # Create new Document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading(project_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_run.italic = True
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph()
        
        # Parse and add markdown content
        self._add_markdown_to_docx(doc, documentation_content)
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Create filename
        filename = f"{self._sanitize_filename(project_name)}_documentation.docx"
        
        return file_stream.read(), filename
    
    def export_pdf(self, project_name, documentation_content):
        """
        Export documentation as PDF file.
        
        Args:
            project_name: Name of the project
            documentation_content: Markdown content string
            
        Returns:
            tuple: (file_bytes, filename)
            
        Raises:
            RuntimeError: If WeasyPrint is not available
        """
        # Check if WeasyPrint is available
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError(
                "PDF export is not available. WeasyPrint requires GTK libraries. "
                "On Windows, you can download GTK from: https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases "
                "Or use Markdown (.md) or DOCX (.docx) export instead."
            )
        
        # Convert markdown to HTML
        html_content = markdown2.markdown(
            documentation_content,
            extras=['fenced-code-blocks', 'tables', 'header-ids']
        )
        
        # Create full HTML document with styling
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{project_name} - Documentation</title>
        </head>
        <body>
            <div class="header">
                <h1>{project_name}</h1>
                <p class="date">Generated on {datetime.now().strftime('%B %d, %Y')}</p>
                <hr>
            </div>
            <div class="content">
                {html_content}
            </div>
        </body>
        </html>
        """
        
        # CSS styling for PDF
        css = CSS(string="""
            @page {
                size: A4;
                margin: 2cm;
                @top-center {
                    content: string(project-name);
                    font-size: 10pt;
                    color: #666;
                }
                @bottom-center {
                    content: "Page " counter(page) " of " counter(pages);
                    font-size: 10pt;
                    color: #666;
                }
            }
            
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 100%;
            }
            
            .header {
                text-align: center;
                margin-bottom: 30px;
            }
            
            .header h1 {
                font-size: 32pt;
                color: #1e40af;
                margin-bottom: 10px;
                string-set: project-name content();
            }
            
            .header .date {
                font-size: 11pt;
                color: #666;
                font-style: italic;
            }
            
            .header hr {
                border: none;
                border-top: 2px solid #3b82f6;
                margin: 20px 0;
            }
            
            .content h1 {
                font-size: 24pt;
                color: #1e40af;
                margin-top: 30px;
                margin-bottom: 15px;
                border-bottom: 2px solid #3b82f6;
                padding-bottom: 5px;
            }
            
            .content h2 {
                font-size: 20pt;
                color: #1e40af;
                margin-top: 25px;
                margin-bottom: 12px;
            }
            
            .content h3 {
                font-size: 16pt;
                color: #2563eb;
                margin-top: 20px;
                margin-bottom: 10px;
            }
            
            .content h4 {
                font-size: 14pt;
                color: #2563eb;
                margin-top: 15px;
                margin-bottom: 8px;
            }
            
            .content p {
                margin-bottom: 12px;
                text-align: justify;
            }
            
            .content ul, .content ol {
                margin-left: 25px;
                margin-bottom: 15px;
            }
            
            .content li {
                margin-bottom: 5px;
            }
            
            .content code {
                background-color: #f3f4f6;
                padding: 2px 6px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                font-size: 10pt;
                color: #dc2626;
            }
            
            .content pre {
                background-color: #1e293b;
                color: #e2e8f0;
                padding: 15px;
                border-radius: 5px;
                overflow-x: auto;
                margin-bottom: 15px;
                font-family: 'Courier New', monospace;
                font-size: 9pt;
                line-height: 1.4;
            }
            
            .content pre code {
                background-color: transparent;
                padding: 0;
                color: inherit;
            }
            
            .content blockquote {
                border-left: 4px solid #3b82f6;
                padding-left: 15px;
                margin-left: 0;
                color: #666;
                font-style: italic;
            }
            
            .content table {
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 20px;
            }
            
            .content table th {
                background-color: #3b82f6;
                color: white;
                padding: 10px;
                text-align: left;
                font-weight: bold;
            }
            
            .content table td {
                border: 1px solid #ddd;
                padding: 8px;
            }
            
            .content table tr:nth-child(even) {
                background-color: #f9fafb;
            }
            
            .content a {
                color: #3b82f6;
                text-decoration: none;
            }
            
            .content a:hover {
                text-decoration: underline;
            }
            
            /* Page break control */
            h1, h2, h3 {
                page-break-after: avoid;
            }
            
            pre, table {
                page-break-inside: avoid;
            }
        """)
        
        # Generate PDF
        pdf_bytes = HTML(string=full_html).write_pdf(stylesheets=[css])
        
        # Create filename
        filename = f"{self._sanitize_filename(project_name)}_documentation.pdf"
        
        return pdf_bytes, filename
    
    def _add_markdown_to_docx(self, doc, markdown_content):
        """
        Parse markdown content and add to DOCX document.
        
        Args:
            doc: python-docx Document object
            markdown_content: Markdown string
        """
        lines = markdown_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Headings (## Heading)
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                
                if level == 1:
                    heading = doc.add_heading(text, level=1)
                elif level == 2:
                    heading = doc.add_heading(text, level=2)
                elif level == 3:
                    heading = doc.add_heading(text, level=3)
                else:
                    heading = doc.add_heading(text, level=4)
                
                i += 1
                continue
            
            # Bullet lists (- item or * item)
            if line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
                # Collect all list items
                list_items = []
                while i < len(lines):
                    current_line = lines[i].rstrip()
                    if current_line.startswith('- ') or current_line.startswith('* ') or current_line.startswith('• '):
                        list_items.append(current_line[2:].strip())
                        i += 1
                    elif not current_line:
                        i += 1
                        break
                    else:
                        break
                
                # Add list items
                for item in list_items:
                    para = doc.add_paragraph(style='List Bullet')
                    self._add_formatted_text(para, item)
                
                continue
            
            # Numbered lists (1. item)
            if re.match(r'^\d+\.\s', line):
                # Collect all list items
                list_items = []
                while i < len(lines):
                    current_line = lines[i].rstrip()
                    if re.match(r'^\d+\.\s', current_line):
                        list_items.append(re.sub(r'^\d+\.\s', '', current_line))
                        i += 1
                    elif not current_line:
                        i += 1
                        break
                    else:
                        break
                
                # Add list items
                for item in list_items:
                    para = doc.add_paragraph(style='List Number')
                    self._add_formatted_text(para, item)
                
                continue
            
            # Regular paragraphs
            para = doc.add_paragraph()
            self._add_formatted_text(para, line)
            i += 1
    
    def _add_formatted_text(self, paragraph, text):
        """
        Add text with inline formatting (bold, italic, code) to paragraph.
        
        Args:
            paragraph: python-docx Paragraph object
            text: Text string with markdown formatting
        """
        # Parse inline formatting
        parts = []
        current = ""
        i = 0
        
        while i < len(text):
            # Bold (**text**)
            if text[i:i+2] == '**':
                if current:
                    parts.append(('normal', current))
                    current = ""
                
                # Find closing **
                end = text.find('**', i + 2)
                if end != -1:
                    parts.append(('bold', text[i+2:end]))
                    i = end + 2
                else:
                    current += text[i]
                    i += 1
            
            # Code (`code`)
            elif text[i] == '`':
                if current:
                    parts.append(('normal', current))
                    current = ""
                
                # Find closing `
                end = text.find('`', i + 1)
                if end != -1:
                    parts.append(('code', text[i+1:end]))
                    i = end + 1
                else:
                    current += text[i]
                    i += 1
            
            # Italic (*text*)
            elif text[i] == '*':
                if current:
                    parts.append(('normal', current))
                    current = ""
                
                # Find closing *
                end = text.find('*', i + 1)
                if end != -1:
                    parts.append(('italic', text[i+1:end]))
                    i = end + 1
                else:
                    current += text[i]
                    i += 1
            
            else:
                current += text[i]
                i += 1
        
        if current:
            parts.append(('normal', current))
        
        # Add runs with formatting
        for style, content in parts:
            run = paragraph.add_run(content)
            
            if style == 'bold':
                run.bold = True
            elif style == 'italic':
                run.italic = True
            elif style == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(220, 38, 38)
    
    def _sanitize_filename(self, name):
        """
        Sanitize filename by removing invalid characters.
        
        Args:
            name: Original filename
            
        Returns:
            str: Sanitized filename
        """
        # Remove or replace invalid characters
        name = re.sub(r'[<>:"/\\|?*]', '', name)
        name = name.strip()
        name = name.replace(' ', '_')
        
        return name or 'documentation'

